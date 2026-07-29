"""
Costruttore DOCX per le recensioni bookmaker Bottadiculo.
Replica esatta dello stile del template Recensione SNAI 2026 (approvato),
estratto direttamente dal file docx originale (colori, font, tabelle, box).

Uso: vedi build_betflag_docx.py per un esempio completo di utilizzo.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Palette esatta (estratta da Recensione_SNAI_2026.docx) ---
NAVY = "1A1A2E"        # titoli H1, sfondo header tabelle dati
BLUE = "0F3460"        # H2, voto numerico, box "nota importante"
GOLD = "C8A951"        # barra CTA, stella voto, divisori, voto medio, box voto finale
RED = "C0392B"         # box "perché non fare al caso tuo"
BODY = "333333"        # testo corpo
GRAY = "666666"        # header/footer pagina, sottotitoli box rating
BULLET_GRAY = "555555"
BORDER = "CCCCCC"
ROW_ALT = "F5F5F5"
ROW_WHITE = "FFFFFF"
PINK_FILL = "FDF2F2"
CREAM_FILL = "FFF9E6"
LIGHTBLUE_FILL = "EBF5FB"
FOOTER_BORDER = "E0E0E0"


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color=BORDER, sz=1):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_paragraph_shading(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_paragraph_bottom_border(paragraph, color=GOLD, sz=6, space=1):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(paragraph, text, bold=False, italic=False, color=BODY, size=10, font=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if font:
        run.font.name = font
    return run


def _set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def new_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    # stile base
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    return doc


def set_header_footer(doc, header_text):
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    hdr_p = section.header.paragraphs[0]
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_bottom_border(hdr_p, color=GOLD, sz=2, space=4)
    _add_run(hdr_p, header_text, italic=True, color=GRAY, size=8)

    ftr_p = section.footer.paragraphs[0]
    ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = ftr_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "1")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), FOOTER_BORDER)
    pBdr.append(top)
    pPr.append(pBdr)

    _add_run(ftr_p, "Pagina ", color=GRAY, size=8)
    run = ftr_p.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    _add_run(ftr_p, " | Contenuto editoriale indipendente", color=GRAY, size=8)


def add_title(doc, title_text, subtitle_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    _add_run(p, title_text, bold=True, color=NAVY, size=20)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(3)
    _add_run(p2, subtitle_text, color=BLUE, size=13)


def add_rating_box(doc, rating_text, col2_title, col2_sub, col3_title, col3_sub):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    widths = (Cm(5.3), Cm(5.3), Cm(5.3))
    cells = table.rows[0].cells
    for cell, w in zip(cells, widths):
        cell.width = w
        _set_cell_margins(cell)

    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, rating_text, bold=True, color=GOLD, size=14)

    p = cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, col2_title, bold=True, color=NAVY, size=11)
    p2 = cells[1].add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, col2_sub, color=GRAY, size=9)

    p = cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, col3_title, bold=True, color=NAVY, size=11)
    p2 = cells[2].add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, col3_sub, color=GRAY, size=9)
    return table


def add_cta_bar(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    _set_paragraph_shading(p, GOLD)
    _add_run(p, f"  {text}  ", bold=True, color="FFFFFF", size=11)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    _set_paragraph_bottom_border(p, color=GOLD, sz=6, space=1)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    _add_run(p, text, bold=True, color=NAVY, size=14)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    _add_run(p, text, bold=True, color=BLUE, size=12)
    return p


def add_body(doc, text, size=10, color=BODY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    # supporta **grassetto** inline in stile markdown semplice
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        _add_run(p, part, bold=(i % 2 == 1), color=color, size=size)
    return p


def add_info_table(doc, rows):
    """rows: lista di tuple (label, value)"""
    table = doc.add_table(rows=0, cols=2)
    col_widths = (Cm(4.5), Cm(11.0))
    for i, (label, value) in enumerate(rows):
        row_cells = table.add_row().cells
        fill = ROW_ALT if i % 2 == 0 else ROW_WHITE
        for cell, w in zip(row_cells, col_widths):
            cell.width = w
            _set_cell_shading(cell, fill)
            _set_cell_borders(cell)
            _set_cell_margins(cell)
        p0 = row_cells[0].paragraphs[0]
        _add_run(p0, label, bold=True, color=BODY, size=10)
        p1 = row_cells[1].paragraphs[0]
        _add_run(p1, value, color=BODY, size=10)
    return table


def add_callout_box(doc, heading, bullets, border_color=RED, fill=PINK_FILL, text_color=BULLET_GRAY):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Cm(15.5)
    _set_cell_shading(cell, fill)
    _set_cell_borders(cell, color=border_color, sz=3)
    _set_cell_margins(cell, top=120, left=200, bottom=120, right=200)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, heading, bold=True, color=border_color, size=10)

    for b in bullets:
        bp = cell.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(2)
        _add_run(bp, b, color=text_color, size=9.5)
    return table


def add_note_box(doc, label, text, border_color=BLUE, fill=LIGHTBLUE_FILL):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Cm(15.5)
    _set_cell_shading(cell, fill)
    _set_cell_borders(cell, color=border_color, sz=2)
    _set_cell_margins(cell, top=120, left=200, bottom=120, right=200)
    p = cell.paragraphs[0]
    _add_run(p, f"{label}: ", bold=True, color=border_color, size=10)
    _add_run(p, text, color=BODY, size=10)
    return table


def add_data_table(doc, headers, rows, highlight_col=None, highlight_color=BLUE, col_widths=None):
    """Tabella dati generica (Valutazioni, Tabella comparativa).
    headers: lista di stringhe
    rows: lista di liste (stessa lunghezza di headers)
    highlight_col: indice colonna da colorare in blu e centrare (es. Voto)
    """
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    if col_widths is None:
        col_widths = [Cm(15.5 / ncols)] * ncols

    header_cells = table.rows[0].cells
    for i, (cell, h, w) in enumerate(zip(header_cells, headers, col_widths)):
        cell.width = w
        _set_cell_shading(cell, NAVY)
        _set_cell_borders(cell)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, h, bold=True, color="FFFFFF", size=10)

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = ROW_ALT if ridx % 2 == 0 else ROW_WHITE
        for cidx, (cell, val, w) in enumerate(zip(cells, row, col_widths)):
            cell.width = w
            _set_cell_shading(cell, fill)
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            if cidx == highlight_col:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, val, bold=True, color=highlight_color, size=10)
            elif cidx == 0:
                _add_run(p, val, bold=True, color=BODY, size=10)
            else:
                _add_run(p, val, color=BODY, size=10)
    return table


def add_average_score(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    _add_run(p, text, bold=True, color=GOLD, size=11)


def add_final_score_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Cm(15.5)
    _set_cell_shading(cell, CREAM_FILL)
    _set_cell_borders(cell, color=GOLD, sz=4)
    _set_cell_margins(cell, top=160, left=200, bottom=160, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text, bold=True, color=NAVY, size=16)
    return table


def add_bullets(doc, items, color=BODY, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        parts = it.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            _add_run(p, part, bold=(i % 2 == 1), color=color, size=size)
